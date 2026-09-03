import re
from copy import deepcopy
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import load_workbook


COA_FIELDS = [
    ("date", "Date", True), ("product_name", "Name of Product", True),
    ("product_code", "Product Code", True), ("batch_number", "Batch Number", True),
    ("customer_name", "Customer Name", False), ("manufacturing_date", "Date of Manufacturing", True),
    ("expiry_date", "Expiry Date", True), ("quantity", "Quantity", False),
    ("storage_condition", "Storage Condition", False), ("tested_by", "Tested By", False),
    ("checked_by", "Checked By", False),
]

SDS_FIELDS = [
    ("product_identifier", "Product identifier", True), ("other_identifiers", "Other identifiers", False),
    ("recommended_use", "Recommended use", True), ("supplier_name", "Supplier / company name", True),
    ("supplier_address", "Supplier address", True), ("supplier_contact", "Supplier contact", False), ("supplier_phone", "Supplier phone", True),
    ("emergency_phone", "Emergency phone", True), ("classification", "Hazard classification", True),
    ("signal_word", "Signal word", True), ("hazard_statements", "Hazard statements", True),
    ("precautionary_statements", "Precautionary statements", True), ("other_hazards", "Other hazards", False),
    ("appearance", "Appearance", True), ("colour", "Colour", False), ("odour", "Odour", False),
    ("ph", "pH", False), ("melting_point", "Melting / freezing point", False),
    ("boiling_point", "Boiling point / range", False), ("flash_point", "Flash point", True),
    ("flammability", "Flammability", False), ("auto_ignition", "Auto-ignition temperature", False),
    ("decomposition_temperature", "Decomposition temperature", False), ("viscosity", "Viscosity", False),
    ("solubility", "Solubility", False), ("vapour_pressure", "Vapour pressure", False),
    ("relative_density", "Relative density", False), ("revision_date", "Revision date", True),
]


def normalise(value: Any) -> str:
    return re.sub(r"[^a-z0-9]+", " ", str(value or "").lower()).strip()


def _set_text_preserving_first_run(paragraph, value: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)


def _replace_value_after_colon(paragraph, value: str) -> None:
    """Replace only the value while preserving the template's label, tabs, and run formatting."""
    text = paragraph.text
    colon = text.find(":")
    if colon < 0:
        return
    suffix = text[colon + 1:]
    spacing = re.match(r"[\s\t]*", suffix).group(0)
    value_start = colon + 1 + len(spacing)
    cursor = 0
    inserted = False
    for run in paragraph.runs:
        original = run.text
        end = cursor + len(original)
        if end <= value_start:
            pass
        elif cursor <= value_start < end:
            run.text = original[:value_start - cursor] + value
            inserted = True
        else:
            run.text = ""
        cursor = end
    if not inserted:
        target = paragraph.runs[-1] if paragraph.runs else paragraph.add_run()
        target.text += value


def _replace_multiple_labelled_values(paragraph, document_type: str, fields: dict[str, str]) -> set[str]:
    """Fill two or more labels sharing one paragraph without deleting later labels.

    The approved COA keeps Tested By and Checked By in one tab-separated
    paragraph. Replacing everything after the first colon removed Checked By.
    """
    definitions = COA_FIELDS if document_type == "coa" else SDS_FIELDS
    labels = {normalise(label): (label, key) for key, label, _ in definitions}
    pattern = "|".join(re.escape(item[0]) for item in sorted(labels.values(), key=lambda item: len(item[0]), reverse=True))
    matches = list(re.finditer(rf"(?i)\b({pattern})\s*:", paragraph.text))
    if len(matches) < 2:
        return set()
    original = paragraph.text
    rebuilt = original[:matches[0].start()]
    replaced: set[str] = set()
    for index, match in enumerate(matches):
        key = labels[normalise(match.group(1))][1]
        segment_end = matches[index + 1].start() if index + 1 < len(matches) else len(original)
        old_value = original[match.end():segment_end]
        leading = re.match(r"[\s\t]*", old_value).group(0)
        trailing_match = re.search(r"[\s\t]*$", old_value)
        trailing = trailing_match.group(0) if trailing_match else ""
        rebuilt += original[match.start():match.end()] + leading + str(fields.get(key, "")) + trailing
        replaced.add(key)
    _set_text_preserving_first_run(paragraph, rebuilt)
    return replaced


def _field_aliases(document_type: str) -> dict[str, str]:
    fields = COA_FIELDS if document_type == "coa" else SDS_FIELDS
    aliases = {normalise(label): key for key, label, _ in fields}
    aliases.update({normalise(key): key for key, _, _ in fields})
    aliases.update({
        "name of product": "product_name", "name of the product": "product_name", "product name": "product_identifier",
        "product use": "recommended_use", "product uses": "recommended_use", "uses advised against": "recommended_use",
        "company details": "supplier_name", "company name": "supplier_name",
        "address": "supplier_address", "company address": "supplier_address", "contact": "supplier_contact", "contact number": "supplier_phone",
        "company phone": "supplier_phone", "telephone": "supplier_phone", "emergency telephone": "emergency_phone",
        "1 4 emergency telephone number": "emergency_phone", "class and category of danger": "classification",
        "color": "colour", "odor": "odour", "odor odor threshold": "odour", "density": "relative_density", "density at 15 c": "relative_density",
        "boiling range": "boiling_point", "vapour pressure at 20 c": "vapour_pressure",
        "kinematic viscosity": "viscosity", "date of revision": "revision_date",
    })
    return aliases


def _replace_labelled_paragraphs(document, document_type: str, fields: dict[str, str]) -> set[str]:
    aliases = _field_aliases(document_type)
    replaced: set[str] = set()
    multiline_labels = {"company address", "class and category of danger", "hazard statements", "supplemental information", "precautionary statements"}
    continuing = False
    for paragraph in document.paragraphs:
        multiple = _replace_multiple_labelled_values(paragraph, document_type, fields)
        if multiple:
            replaced.update(multiple)
            continuing = False
            continue
        text = paragraph.text.strip()
        if ":" not in text:
            if continuing and text:
                _set_text_preserving_first_run(paragraph, "")
            if not text or text.lower().startswith("section ") or re.match(r"^\d+(?:\.\d+)*\s", text):
                continuing = False
            continue
        label = normalise(text.split(":", 1)[0])
        key = aliases.get(label) or f"template__{label.replace(' ', '_')}"
        _replace_value_after_colon(paragraph, fields.get(key, ""))
        replaced.add(key)
        continuing = label in multiline_labels
    return replaced


def _replace_table_rows(table, rows: list[dict[str, str]], keys: list[str]) -> None:
    template_row = deepcopy(table.rows[1]._tr if len(table.rows) > 1 else table.rows[0]._tr)
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    for row_data in rows:
        new_row = deepcopy(template_row)
        table._tbl.append(new_row)
        row = table.rows[-1]
        for index, key in enumerate(keys):
            if index < len(row.cells):
                row.cells[index].text = str(row_data.get(key, ""))


def coa_parameter_rows(template: Path) -> list[dict[str, str]]:
    document = Document(template)
    if not document.tables:
        return []
    return [{"parameter": row.cells[0].text.strip(), "specification": "", "result": ""} for row in document.tables[0].rows[1:] if row.cells and row.cells[0].text.strip()]


def _fill_coa_rows(table, supplied_rows: list[dict[str, str]]) -> None:
    supplied = {normalise(row.get("parameter")): row for row in supplied_rows if row.get("parameter")}
    for row in table.rows[1:]:
        parameter = row.cells[0].text.strip()
        values = supplied.get(normalise(parameter), {})
        if len(row.cells) > 1:
            row.cells[1].text = str(values.get("specification", ""))
        if len(row.cells) > 2:
            row.cells[2].text = str(values.get("result", ""))


def generate_docx(template: Path, output: Path, document_type: str, fields: dict[str, str], rows: list[dict[str, str]]) -> list[str]:
    document = Document(template)
    warnings: list[str] = []
    rendered_fields = dict(fields)
    if document_type == "coa" and rendered_fields.get("date"):
        try:
            rendered_fields["date"] = datetime.strptime(rendered_fields["date"], "%d %B %Y").strftime("%d-%m-%Y")
        except ValueError:
            pass
    replaced = _replace_labelled_paragraphs(document, document_type, rendered_fields)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replaced.update(_replace_labelled_paragraphs(type("Part", (), {"paragraphs": [paragraph]})(), document_type, rendered_fields))
    definitions = COA_FIELDS if document_type == "coa" else SDS_FIELDS
    missing = [label for key, label, required in definitions if required and not str(fields.get(key, "")).strip()]
    if missing:
        warnings.append("Missing required information: " + ", ".join(missing))
    if document_type == "coa":
        if document.tables:
            _fill_coa_rows(document.tables[0], rows)
        if not any(row.get("specification") or row.get("result") for row in rows):
            warnings.append("COA test parameters were retained, but their specification and result values are blank.")
    else:
        if document.tables:
            _replace_table_rows(document.tables[0], rows, ["name", "cas_number", "ec_number", "concentration", "classification", "notes"])
            # Clear other product-specific data tables, while retaining the final abbreviation/reference table.
            for table in document.tables[1:-1]:
                for row in table.rows[1:]:
                    for cell in row.cells:
                        cell.text = ""
        if not rows:
            warnings.append("No SDS composition rows were supplied; the composition table is blank.")
        warnings.append("SDS documents require review and approval by a qualified safety/regulatory person before issue.")
    if not replaced:
        warnings.append("No labelled fields were found in this template; verify the generated document carefully.")
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return warnings


def read_excel(path: Path, document_type: str) -> tuple[dict[str, str], list[dict[str, str]]]:
    workbook = load_workbook(path, read_only=True, data_only=True)
    fields: dict[str, str] = {}
    rows: list[dict[str, str]] = []
    definitions = COA_FIELDS if document_type == "coa" else SDS_FIELDS
    aliases = _field_aliases(document_type)
    row_keys = ["parameter", "specification", "result"] if document_type == "coa" else ["name", "cas_number", "ec_number", "concentration", "classification", "notes"]
    for sheet in workbook.worksheets:
        values = list(sheet.iter_rows(values_only=True))
        if not values:
            continue
        title = normalise(sheet.title)
        if title in {"parameters", "composition"}:
            headers = [normalise(value).replace(" ", "_") for value in values[0]]
            for values_row in values[1:]:
                record = {headers[i]: str(value) for i, value in enumerate(values_row) if i < len(headers) and value is not None}
                canonical = {key: record.get(key, "") for key in row_keys}
                if any(canonical.values()):
                    rows.append(canonical)
            continue
        # Accept either Field/Value columns or a simple two-column sheet.
        start = 1 if normalise(values[0][0] if values[0] else "") in {"field", "label", "name"} else 0
        for values_row in values[start:]:
            if len(values_row) < 2 or values_row[0] is None:
                continue
            key = aliases.get(normalise(values_row[0]))
            if key and values_row[1] is not None:
                fields[key] = str(values_row[1])
    workbook.close()
    return fields, rows


def field_schema(document_type: str, template: Path | None = None) -> list[dict[str, Any]]:
    definitions = COA_FIELDS if document_type == "coa" else SDS_FIELDS
    result = [{"key": key, "label": label, "required": required} for key, label, required in definitions]
    if document_type != "sds" or not template:
        return result
    aliases = _field_aliases(document_type)
    existing = {item["key"] for item in result}
    for paragraph in Document(template).paragraphs:
        text = paragraph.text.strip()
        if ":" not in text:
            continue
        shown_label = text.split(":", 1)[0].strip()
        label = normalise(shown_label)
        key = aliases.get(label) or f"template__{label.replace(' ', '_')}"
        if key not in existing:
            result.append({"key": key, "label": shown_label, "required": False})
            existing.add(key)
    return result
