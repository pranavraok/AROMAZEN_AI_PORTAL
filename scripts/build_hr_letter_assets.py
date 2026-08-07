from __future__ import annotations

import shutil
import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "backend" / "app" / "assets" / "hr_letters"
DOWNLOADS = Path(r"C:\Users\prana\Downloads")


def replace_paragraph(paragraph, replacements: dict[str, str]) -> None:
    if not paragraph.runs:
        return
    for old, new in replacements.items():
        while old in "".join(run.text for run in paragraph.runs):
            combined = "".join(run.text for run in paragraph.runs)
            start = combined.index(old)
            end = start + len(old)
            offsets: list[tuple[int, int]] = []
            cursor = 0
            for run in paragraph.runs:
                offsets.append((cursor, cursor + len(run.text)))
                cursor += len(run.text)
            start_run = next(i for i, (_, right) in enumerate(offsets) if right > start)
            end_run = next(i for i, (_, right) in enumerate(offsets) if right >= end)
            left_start, _ = offsets[start_run]
            left_end, _ = offsets[end_run]
            prefix = paragraph.runs[start_run].text[: start - left_start]
            suffix = paragraph.runs[end_run].text[end - left_end :]
            paragraph.runs[start_run].text = prefix + new + suffix
            for index in range(start_run + 1, end_run + 1):
                paragraph.runs[index].text = ""


def all_paragraphs(document: Document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
                for nested in cell.tables:
                    for nested_row in nested.rows:
                        for nested_cell in nested_row.cells:
                            yield from nested_cell.paragraphs
    for section in document.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def apply(document: Document, replacements: dict[str, str]) -> None:
    for paragraph in all_paragraphs(document):
        replace_paragraph(paragraph, replacements)


def salary_placeholders(document: Document) -> None:
    rows = []
    for table in document.tables:
        rows.extend(table.rows)
        for row in table.rows:
            for cell in row.cells:
                for nested in cell.tables:
                    rows.extend(nested.rows)
    for row_index, row in enumerate(rows):
        if len(row.cells) < 3 or row_index < 2:
            continue
        label = row.cells[0].text.strip().lower()
        if not label or label == "salary components" or label.startswith("*") or label.startswith("note:"):
            continue
        key = re.sub(r"[^a-z0-9]+", "_", label).strip("_")[:48]
        for column, suffix in ((1, "existing"), (2, "revised")):
            paragraph = row.cells[column].paragraphs[0]
            if paragraph.text.strip():
                replace_paragraph(paragraph, {paragraph.text: f"{{{{salary_{key}_{suffix}}}}}"})


def build_special_increment() -> None:
    document = Document(DOWNLOADS / "Special increment letter.docx")
    apply(document, {
        "05/08/2026": "{{issue_date}}",
        "Sanjana Sadananda Manjrekar": "{{employee_name}}",
        "05th August 2026": "{{effective_date_long}}",
        "₹21,608/-": "{{revised_gross_salary}}",
        "APL/MG/009/25-26": "{{appointment_reference}}",
        "16/06/2027": "{{appointment_date}}",
        "Ms. Swathi Nayak": "{{signatory_name}}",
        "Human Resources": "{{signatory_title}}",
    })
    salary_placeholders(document)
    if document.paragraphs and "confidential information" in document.paragraphs[-1].text.lower():
        document.paragraphs[-1]._element.getparent().remove(document.paragraphs[-1]._element)
    document.save(OUT / "special-increment-template.docx")


def build_appointment() -> None:
    document = Document(DOWNLOADS / "appointment letter.docx")
    apply(document, {
        "APL/MG/006/26-27": "{{reference_number}}",
        "Ms.Deeksha Shettigar": "{{employee_salutation_name}}",
        "Deeksha Shettigar": "{{employee_name}}",
        "ದೀಕ್ಷಾ ಶೆಟ್ಟಿಗಾರ್": "{{employee_name_kannada}}",
        "#3-131/1 Tharapalke House\nTalipady, Kinnigoli\nMangalore, Dakshina Kannada \nKarnataka - 574150": "{{employee_address}}",
        "97422560579": "{{employee_phone}}",
        "Sourcing Executive": "{{designation}}",
        "ಸೋರ್ಸಿಂಗ್ ಕಾರ್ಯನಿರ್ವಾಹಕಿ": "{{designation_kannada}}",
        "Mrs. Deeksha": "{{reporting_officer}}",
        "ಶ್ರೀಮತಿ ದೀಕ್ಷಾ": "{{reporting_officer_kannada}}",
        "33,500/-": "{{gross_salary}}",
        "Thirty-Three Thousand Five Hundred": "{{gross_salary_words}}",
        "ಮೂವತ್ತಮೂರು ಸಾವಿರ_ ಐದುನೂರು": "{{gross_salary_words_kannada}}",
        "Ms.Swathi Nayak": "{{signatory_name}}",
        "ಸ್ವಾತಿ ನಾಯಕ್": "{{signatory_name_kannada}}",
        "Human Resources": "{{signatory_title}}",
    })
    # The first date is the issue date; remaining occurrences are the joining/acceptance date.
    first_date_replaced = False
    for paragraph in all_paragraphs(document):
        if "01/06/2026" not in paragraph.text:
            continue
        replacement = "{{issue_date}}" if not first_date_replaced else "{{joining_date}}"
        replace_paragraph(paragraph, {"01/06/2026": replacement})
        first_date_replaced = True
    salary_placeholders(document)
    document.save(OUT / "appointment-template.docx")


def build_spot_appreciation() -> None:
    document = Document(DOWNLOADS / "SPOT APPRECIATION LETTER.docx")
    paragraphs = document.paragraphs
    for paragraph in list(paragraphs[22:]):
        paragraph._element.getparent().remove(paragraph._element)
    apply(document, {
        "24-06-2026": "{{issue_date}}",
        "Mr. Manoj Kumar": "{{employee_name}}",
        "Production Assistant": "{{designation}}",
        "It gives us great pleasure to appreciate your alertness and proactive approach in identifying the incorrect material before it was used in production.": "{{appreciation_reason}}",
        "Your timely observation helped prevent production issues, avoided material wastage, and ensure smooth operations. This reflects your attention to detail, commitment to quality, and sense of responsibility towards our manufacturing processes.": "{{positive_impact}}",
        "We recognise your initiative and the positive impact of your contribution. As a token of our appreciation, please accept a cash reward of ₹5,000/-.": "{{reward_statement}}",
        "Thank you for your dedication and good work. Keep up the excellent performance.": "{{closing_message}}",
        "Mrs. Deeksha": "{{signatory_name}}",
        "Accounts & HR Head": "{{signatory_title}}",
    })
    document.save(OUT / "spot-appreciation-template.docx")


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    build_special_increment()
    build_appointment()
    build_spot_appreciation()
    shutil.copy2(DOWNLOADS / "offer letter template.pdf", OUT / "offer-template.pdf")


if __name__ == "__main__":
    main()
