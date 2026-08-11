from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document


SALARY_COMPONENTS = [
    "basic_salary",
    "house_rent_allowance",
    "conveyance_allowance",
    "medical_reimbursement",
    "special_allowance",
    "variable_pay",
    "a_gross_salary_total",
    "employer_contribution_pf_12",
    "employer_contribution_esi_3_25",
    "b_employee_retiral_total",
    "residential_telephone_reimbursement",
    "car_fuel_and_maintenance_allowance",
    "monthly_perquisites_total",
    "group_medical_insurance_policy_premium_gmc",
    "group_personal_accident_policy_gpa",
    "employee_benefits_total",
    "fixed_compensation_in_hand",
    "cost_to_company_ctc_per_annum_total",
]

def replace_token(paragraph, token: str, value: str) -> int:
    replacements = 0
    while token in "".join(run.text for run in paragraph.runs):
        combined = "".join(run.text for run in paragraph.runs)
        start = combined.index(token)
        end = start + len(token)
        offsets: list[tuple[int, int]] = []
        cursor = 0
        for run in paragraph.runs:
            offsets.append((cursor, cursor + len(run.text)))
            cursor += len(run.text)
        start_run = next(i for i, (_, right) in enumerate(offsets) if right > start)
        end_run = next(i for i, (_, right) in enumerate(offsets) if right >= end)
        start_left, _ = offsets[start_run]
        end_left, _ = offsets[end_run]
        prefix = paragraph.runs[start_run].text[: start - start_left]
        suffix = paragraph.runs[end_run].text[end - end_left :]
        paragraph.runs[start_run].text = prefix + value + suffix
        for index in range(start_run + 1, end_run + 1):
            paragraph.runs[index].text = ""
        replacements += 1
    return replacements


def all_paragraphs(document: Document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in document.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def replace_everywhere(document: Document, token: str, value: str) -> int:
    return sum(replace_token(paragraph, token, value) for paragraph in all_paragraphs(document))


def set_blank_cell_token(cell, token: str) -> None:
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        paragraph.runs[0].text = token
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(token)


def build(source: Path, output: Path) -> None:
    document = Document(source)
    if len(document.tables) != 1 or len(document.tables[0].rows) != 24:
        raise ValueError("The appointment draft must retain its approved 24-row compensation table.")

    paragraphs = document.paragraphs
    replace_everywhere(document, "APL/MG/xxx/26-27", "{{reference_number}}")
    replace_token(paragraphs[4], "(DOJ)", "{{issue_date}}")
    replace_token(paragraphs[6], "Ms. SURAKSHA", "{{employee_salutation_name}}")
    replace_token(paragraphs[7], "(Address)", "{{employee_address}}")
    replace_token(paragraphs[7], "xxxxxxxxxxx", "{{employee_phone}}")
    replace_token(paragraphs[11], "ಆತ್ಮೀಯ ಕು. ", "ಆತ್ಮೀಯ ಕು.{{employee_name_kannada}}")

    # The annexure date is the appointment-letter issue date; other dummy dates
    # in the approved draft represent the employee's joining date.
    replace_token(paragraphs[256], "02/02/2026", "{{issue_date}}")
    replace_everywhere(document, "02/02/2026", "{{joining_date}}")
    replace_token(paragraphs[217], "DOJ", "{{joining_date}}")
    replace_token(paragraphs[219], "DOJ", "{{joining_date}}")

    replace_everywhere(document, "SURAKSHA", "{{employee_name}}")
    replace_everywhere(document, "Lab Assistant", "{{designation}}")
    replace_everywhere(document, "ಪ್ರಯೋಗಾಲಯ ಸಹಾಯಕಿ", "{{designation_kannada}}")
    replace_everywhere(document, "Mrs.Sandhya K", "{{reporting_officer}}")
    replace_everywhere(document, "ಶ್ರೀಮತಿ ಸಂಧ್ಯಾ ಕೆ", "{{reporting_officer_kannada}}")
    replace_everywhere(document, "XXXXX", "{{gross_salary}}")
    replace_everywhere(document, "Thousand_", "{{gross_salary_words}}")
    replace_everywhere(document, "ಹದಿನಾಲ್ಕು ಸಾವಿರ_", "{{gross_salary_words_kannada}}")
    replace_everywhere(document, "Ms.Swathi Nayak", "{{signatory_name}}")
    replace_everywhere(document, "Ms. Swathi Nayak", "{{signatory_name}}")
    replace_everywhere(document, "ಸ್ವಾತಿ ನಾಯಕ್", "{{signatory_name_kannada}}")
    replace_everywhere(document, "Human Resources", "{{signatory_title}}")

    table = document.tables[0]
    for row_index, component in enumerate(SALARY_COMPONENTS, start=2):
        set_blank_cell_token(table.rows[row_index].cells[1], f"{{{{salary_{component}_monthly}}}}")
        set_blank_cell_token(table.rows[row_index].cells[2], f"{{{{salary_{component}_annual}}}}")

    # The two-line designation block is the only approved paragraph that can
    # otherwise split across pages after variable values are inserted.
    paragraphs[29].paragraph_format.keep_together = True

    unresolved_dummy = re.compile(
        r"APL/MG/xxx|SURAKSHA|xxxxxxxxxxx|Lab Assistant|Mrs\.Sandhya|XXXXX|02/02/2026|\bDOJ\b"
    )
    unresolved = [paragraph.text for paragraph in all_paragraphs(document) if unresolved_dummy.search(paragraph.text)]
    if unresolved:
        raise ValueError(f"Unresolved dummy appointment values: {unresolved}")

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> None:
    parser = argparse.ArgumentParser(description="Turn the approved appointment draft into the portal template.")
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build(args.source, args.output)


if __name__ == "__main__":
    main()
