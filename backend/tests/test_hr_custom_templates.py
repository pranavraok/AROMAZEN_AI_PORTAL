from pathlib import Path

from docx import Document
from lxml import etree

from app.modules.hr_letters.routes import (
    _fill_docx,
    _replace_xml_paragraph_tokens,
    _template_tokens,
)


def test_custom_template_detects_and_fills_split_run_placeholder(tmp_path: Path) -> None:
    template = tmp_path / "custom-template.docx"
    document = Document()
    paragraph = document.add_paragraph("Dear ")
    paragraph.add_run("{{ employee")
    paragraph.add_run("_name }}")
    document.save(template)

    assert _template_tokens(template) == ["employee_name"]

    output = _fill_docx(
        "custom",
        {"employee_name": "Pranav Rao"},
        tmp_path,
        source_path=template,
    )
    generated = Document(output)
    assert generated.paragraphs[0].text == "Dear Pranav Rao"
    assert "{{" not in generated.paragraphs[0].text


def test_ooxml_only_placeholder_replacement_handles_multiple_text_nodes() -> None:
    namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    paragraph = etree.fromstring(
        f'<w:p xmlns:w="{namespace}"><w:r><w:t>{{{{job_</w:t></w:r>'
        '<w:r><w:t>title}}</w:t></w:r></w:p>'
    )

    assert _replace_xml_paragraph_tokens(paragraph, {"job_title": "Manager"})
    assert "".join(paragraph.xpath(".//w:t/text()", namespaces={"w": namespace})) == "Manager"
