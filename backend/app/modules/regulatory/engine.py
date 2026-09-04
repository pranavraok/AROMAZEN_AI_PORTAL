from __future__ import annotations

import io
import re
from copy import deepcopy
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import load_workbook


DOCUMENT_TYPES = {"sds", "ifra_certificate", "ifra_amendment", "allergen_report", "reach_declaration"}
INTERNAL_MARKERS = ("ai suggested", "review required", "confidence:", "source url")
OUTPUT_PLACEHOLDERS = {
    "n a",
    "na",
    "not available",
    "no available",
    "no data available",
    "none available",
    "not determined",
    "not applicable",
    "unavailable",
    "unknown",
    "leave blank if unavailable",
}


def normalise(value: Any) -> str:
    return re.sub(r"[^a-z0-9]+", " ", str(value or "").lower()).strip()


def clean_issue_value(value: Any, empty: str = "") -> str:
    text = str(value or "").strip()
    placeholder = normalise(text) in OUTPUT_PLACEHOLDERS
    return empty if placeholder or any(marker in text.lower() for marker in INTERNAL_MARKERS) else text


def parse_regulatory_excel(content: bytes) -> tuple[str, str, list[dict[str, str]]]:
    workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    product = code = ""
    ingredients: list[dict[str, str]] = []
    for sheet in workbook.worksheets:
        rows = list(sheet.iter_rows(values_only=True))
        header_index = -1
        for index, row in enumerate(rows):
            first = normalise(row[0] if row else "")
            if first == "name" and len(row) > 1:
                product = clean_issue_value(row[1])
            elif first == "code" and len(row) > 1:
                code = clean_issue_value(row[1])
            elif first in {"sl no", "serial no", "s no"}:
                header_index = index
        if header_index >= 0:
            for row in rows[header_index + 1:]:
                name = clean_issue_value(row[1] if len(row) > 1 else "")
                concentration = clean_issue_value(row[2] if len(row) > 2 else "")
                if not name or normalise(name) in {"total", "rawmaterials", "raw materials"}:
                    continue
                if not concentration and not any(value is not None for value in row):
                    continue
                ingredients.append({"name": name, "concentration": concentration, "cas": "", "ec": "", "classification": "", "hazard_statements": "", "precautionary_statements": "", "signal_word": "", "pictograms": "", "toxicology": "", "ecology": "", "transport": "", "allergen_identity": "", "svhc_identity": "", "ifra_limits": "", "sources": [], "provenance": "excel"})
            break
    workbook.close()
    if not product or not code or not ingredients:
        raise ValueError("The workbook must contain Name, Code, and an SL NO / RAWMATERIALS / % table.")
    return product, code, ingredients


COA_LABELS = {
    "appearance": ("appearance",), "colour": ("colour", "color"), "odour": ("odour", "odor"),
    "relative_density": ("specific gravity", "relative density", "density"), "flash_point": ("flash point",),
    "refractive_index": ("refractive index",), "solubility": ("solubility",), "storage_condition": ("storage condition", "storage"),
}


def extract_coa_properties(text: str) -> dict[str, str]:
    lines = [" ".join(line.split()) for line in text.splitlines() if line.strip()]
    result: dict[str, str] = {}
    # Word/Excel table extraction uses pipes between cells. Prefer the actual
    # result column, except for qualitative pass/fail rows where the
    # specification is the useful SDS description.
    for line in lines:
        cells = [cell.strip() for cell in line.split("|")]
        if len(cells) < 2:
            continue
        label = normalise(cells[0])
        for key, aliases in COA_LABELS.items():
            if label not in {normalise(alias) for alias in aliases}:
                continue
            values = [cell for cell in cells[1:] if cell]
            if values:
                final = values[-1]
                result[key] = values[0] if normalise(final) in {"pass", "passes", "complies", "conforms"} else final
                if key == "flash_point":
                    result[key] = re.sub(r"(?i)^(\d{1,3})0\s*C$", r"\1°C", result[key])
            break
    all_labels = [alias for aliases in COA_LABELS.values() for alias in aliases]
    all_labels.extend(("tested by", "checked by", "approved by", "parameter", "specification", "result"))
    boundary = "|".join(re.escape(item) for item in sorted(all_labels, key=len, reverse=True))
    for key, aliases in COA_LABELS.items():
        if key in result:
            continue
        for alias in aliases:
            match = re.search(rf"(?i)\b{re.escape(alias)}\b\s*[:|\-]?\s*(.+?)(?=\s+\b(?:{boundary})\b\s*[:|\-]|$)", " ".join(lines))
            if match:
                value = match.group(1).strip(" |,;:-")[:300]
                if value:
                    result[key] = re.sub(r"(?i)^(\d{1,3})0\s*C$", r"\1°C", value) if key == "flash_point" else value
                    break
    return result


def extract_coa_identity(text: str) -> dict[str, str]:
    """Extract only explicit product identifiers from a Creation COA."""
    compact = "\n".join(" ".join(line.split()) for line in text.splitlines() if line.strip())
    result: dict[str, str] = {}
    patterns = {
        "product_name": r"(?im)^\s*(?:name\s+of\s+the\s+product|product\s+name)\s*[:|\-]\s*([^\n|]+)",
        "product_code": r"(?im)^\s*product\s+code\s*[:|\-]\s*([^\n|]+)",
    }
    for key, pattern in patterns.items():
        match = re.search(pattern, compact)
        if match:
            result[key] = clean_issue_value(match.group(1)).strip(" |,;:-")[:300]
    return result


def _set_paragraph_value(paragraph, label_patterns: tuple[str, ...], value: str) -> bool:
    text = paragraph.text
    for pattern in label_patterns:
        match = re.search(rf"(?i)({pattern}\s*:?)(.*)$", text)
        if not match:
            continue
        prefix = text[:match.start(2)]
        cleaned = clean_issue_value(value)
        replacement = prefix.rstrip()
        if cleaned:
            replacement += f" {cleaned}"
        if paragraph.runs:
            paragraph.runs[0].text = replacement
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.text = replacement
        return True
    return False


def _all_paragraphs(document):
    for paragraph in document.paragraphs:
        yield paragraph
    for section in document.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _replace_product(document, product: str, code: str) -> None:
    for paragraph in _all_paragraphs(document):
        _set_paragraph_value(paragraph, (r"^\s*PRODUCT\s+NAME",), product)
        _set_paragraph_value(paragraph, (r"^\s*PRODUCT\s+CODE",), code)
        _set_paragraph_value(paragraph, (r"^\s*Product\s+identifier\s*:",), f"{product} {code}".strip())


def _clear_paragraph(paragraph, remove_drawings: bool = False) -> None:
    for run in paragraph.runs:
        run.text = ""
    if remove_drawings:
        for tag in (".//w:drawing", ".//w:pict"):
            for node in paragraph._p.xpath(tag):
                node.getparent().remove(node)


def _set_paragraph_text_preserving_layout(paragraph, value: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = value


def _remove_output_placeholders(document) -> None:
    """Keep unknown SDS values blank without exposing internal placeholders."""
    suffix = re.compile(
        r"(?i)^(.*?[:：])\s*(?:n\s*/?\s*a|not available|no available|no data available|none available|"
        r"not determined|not applicable|unavailable|unknown)\.?\s*$"
    )
    for paragraph in _all_paragraphs(document):
        text = paragraph.text.strip()
        if not text:
            continue
        if normalise(text) in OUTPUT_PLACEHOLDERS:
            _clear_paragraph(paragraph)
            continue
        match = suffix.match(text)
        if match:
            _set_paragraph_text_preserving_layout(paragraph, match.group(1).rstrip())


def _set_sds_footer_metadata(document, fields: dict[str, str]) -> None:
    version = clean_issue_value(fields.get("version"))
    revision_date = clean_issue_value(fields.get("revision_date"))
    pattern = re.compile(r"(?is)^(.*?Version\s*:)\s*.*?(\s+Date\s*:)\s*.*?$")
    for section in document.sections:
        for paragraph in section.footer.paragraphs:
            match = pattern.match(paragraph.text)
            if not match:
                continue
            replacement = match.group(1)
            if version:
                replacement += f" {version}"
            replacement += match.group(2)
            if revision_date:
                replacement += f" {revision_date}"
            _set_paragraph_text_preserving_layout(paragraph, replacement)


def _remove_duplicate_particle_characteristics(document) -> None:
    found = False
    for paragraph in document.paragraphs:
        if normalise(paragraph.text) != "particle characteristics":
            continue
        if found:
            _clear_paragraph(paragraph)
        found = True


def _paragraph_index(paragraphs: list, pattern: str, start: int = 0) -> int | None:
    regex = re.compile(pattern, re.IGNORECASE)
    return next((index for index in range(start, len(paragraphs)) if regex.search(paragraphs[index].text.strip())), None)


def _replace_labeled_block(document, start_pattern: str, end_pattern: str, value: str, *, remove_drawings: bool = False) -> None:
    paragraphs = list(document.paragraphs)
    start = _paragraph_index(paragraphs, start_pattern)
    if start is None:
        return
    end = _paragraph_index(paragraphs, end_pattern, start + 1)
    if end is None:
        return
    cleaned = clean_issue_value(value)
    if cleaned:
        _set_paragraph_value(paragraphs[start], (start_pattern,), cleaned)
    else:
        _clear_paragraph(paragraphs[start], remove_drawings=remove_drawings)
    for paragraph in paragraphs[start + 1:end]:
        _clear_paragraph(paragraph, remove_drawings=remove_drawings)


def _replace_section_body(document, heading_pattern: str, next_heading_pattern: str, value: str = "") -> None:
    paragraphs = list(document.paragraphs)
    heading = _paragraph_index(paragraphs, heading_pattern)
    if heading is None:
        return
    end = _paragraph_index(paragraphs, next_heading_pattern, heading + 1)
    if end is None:
        return
    _set_paragraph_value(paragraphs[heading], (heading_pattern,), "")
    for offset, paragraph in enumerate(paragraphs[heading + 1:end]):
        if offset == 0 and clean_issue_value(value):
            if paragraph.runs:
                paragraph.runs[0].text = clean_issue_value(value)
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.text = clean_issue_value(value)
        else:
            _clear_paragraph(paragraph, remove_drawings=True)


def _compact_classification(value: Any) -> str:
    text = clean_issue_value(value)
    if not text:
        return ""
    replacements = (
        (r"Specific target organ toxicity\s*-\s*Single exposure", "STOT SE"),
        (r"Specific target organ toxicity\s*-\s*Repeated exposure", "STOT RE"),
        (r"Hazardous to the aquatic environment\s*,?\s*Short term\s*\(Acute\)", "Aquatic Acute"),
        (r"Hazardous to the aquatic environment\s*,?\s*long[- ]term hazard", "Aquatic Chronic"),
        (r"Hazardous to the aquatic environment\s*,?\s*Long term\s*\(Chronic\)", "Aquatic Chronic"),
        (r"Serious eye damage/eye irritation", "Eye Irrit."),
        (r"Skin corrosion/irritation", "Skin Irrit."),
        (r"Sensitization,\s*Skin", "Skin Sens."),
        (r"Flammable liquids", "Flam. Liq."),
        (r"Acute toxicity,\s*oral", "Acute Tox. (oral)"),
        (r"\s*:\s*Category\s*", " "),
    )
    for pattern, replacement in replacements:
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
    text = re.sub(r"\s*\((H\d{3})\)", r": \1", text, flags=re.IGNORECASE)
    text = re.sub(r"\s*;\s*", "; ", text)
    return re.sub(r"\s+", " ", text).strip()


def _format_concentration(value: Any) -> str:
    text = clean_issue_value(value)
    if not text or "%" in text:
        return text
    try:
        number = float(text)
    except ValueError:
        return text
    return f"{number:.1f}%"


def _set_row_no_split(row, *, repeat_header: bool = False) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    for height in list(tr_pr.findall(qn("w:trHeight"))):
        tr_pr.remove(height)
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))
    if repeat_header and tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def _set_compact_cell_text(cell, value: Any, size: float = 8.5) -> None:
    paragraphs = list(cell.paragraphs)
    paragraph = paragraphs[0]
    for extra in paragraphs[1:]:
        cell._tc.remove(extra._p)
    if paragraph.runs:
        run = paragraph.runs[0]
        run.text = clean_issue_value(value)
        for extra in paragraph.runs[1:]:
            extra.text = ""
    else:
        run = paragraph.add_run(clean_issue_value(value))
    run.font.size = Pt(size)
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial")
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _populate_composition_table(table, template_row, ingredients: list[dict]) -> None:
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    _set_row_no_split(table.rows[0], repeat_header=True)
    for item in ingredients:
        table._tbl.append(deepcopy(template_row))
        row = table.rows[-1]
        _set_row_no_split(row)
        values = [
            item.get("name"), item.get("cas"), item.get("ec"),
            _format_concentration(item.get("concentration")),
            _compact_classification(item.get("classification")),
            item.get("specific_concentration_limits"),
        ]
        for index, value in enumerate(values):
            if index < len(row.cells):
                _set_compact_cell_text(row.cells[index], value)


def _ingredient_indexes(ingredients: list[dict[str, Any]]) -> tuple[dict[str, dict], dict[str, dict]]:
    names: dict[str, dict] = {}
    cases: dict[str, dict] = {}
    for item in ingredients:
        for candidate in (item.get("name"), item.get("canonical_name"), item.get("allergen_identity"), item.get("svhc_identity")):
            if normalise(candidate):
                names[normalise(candidate)] = item
        for alias in item.get("aliases") or []:
            if normalise(alias):
                names[normalise(alias)] = item
        for cas in re.split(r"[/,;\s]+", str(item.get("cas") or "")):
            if re.fullmatch(r"\d{2,7}-\d{2}-\d", cas):
                cases[cas] = item
    return names, cases


def _matching_ingredient(name: str, cas_text: str, names: dict[str, dict], cases: dict[str, dict]) -> dict | None:
    normalized_name = normalise(name)
    exact = names.get(normalized_name)
    if exact:
        return exact
    for candidate, item in sorted(names.items(), key=lambda pair: len(pair[0]), reverse=True):
        if len(candidate) >= 4 and re.search(rf"(?:^|\s){re.escape(candidate)}(?:\s|$)", normalized_name):
            return item
    for cas in re.findall(r"\d{2,7}-\d{2}-\d", cas_text):
        if cas in cases:
            return cases[cas]
    return None


def _fill_catalog_tables(document, ingredients: list[dict], absent: str) -> None:
    names, cases = _ingredient_indexes(ingredients)
    for table in document.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            name = row.cells[0].text.strip()
            if normalise(name) in {"ingredient", "name of ingredient", "substance name"}:
                continue
            cas_text = " ".join(cell.text for cell in row.cells[1:-1])
            item = _matching_ingredient(name, cas_text, names, cases)
            target = row.cells[-1]
            _set_cell_text_preserving_layout(
                target,
                clean_issue_value(item.get("concentration"), absent) if item else absent,
            )


def _set_cell_text_preserving_layout(cell, value: Any) -> None:
    paragraphs = list(cell.paragraphs)
    paragraph = paragraphs[0]
    for extra in paragraphs[1:]:
        cell._tc.remove(extra._p)
    if paragraph.runs:
        run = paragraph.runs[0]
        run.text = clean_issue_value(value)
        for extra in paragraph.runs[1:]:
            extra.text = ""
    else:
        paragraph.add_run(clean_issue_value(value))


def _tighten_ifra_certificate(document) -> None:
    if len(document.tables) < 2:
        return
    annex = max(document.tables, key=lambda table: len(table.rows))
    for row_index, row in enumerate(annex.rows):
        _set_row_no_split(row, repeat_header=row_index == 0)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 0.9
                for run in paragraph.runs:
                    if run.font.size is None or run.font.size.pt > 7:
                        run.font.size = Pt(7)


def _fill_sds(document, fields: dict[str, str], ingredients: list[dict]) -> None:
    _set_sds_footer_metadata(document, fields)
    labels = {
        "appearance": (r"Appearance",), "colour": (r"Colou?r",), "odour": (r"Odou?r(?:/Odor threshold)?",),
        "flash_point": (r"Flash point",), "refractive_index": (r"Refractive index",), "solubility": (r"Solubility",),
        "relative_density": (r"Density and/or relative density", r"Relative density"),
        "signal_word": (r"Signal word",), "hazard_statements": (r"Hazard statements",),
        "precautionary_statements": (r"Precautionary statements",),
        "other_hazards": (r"Other hazards", r"Supplemental Information"),
    }
    for paragraph in _all_paragraphs(document):
        for key, patterns in labels.items():
            _set_paragraph_value(paragraph, patterns, clean_issue_value(fields.get(key)))

    # Remove every example-product hazard statement. Mixture classification
    # requires employee review; blank reviewed fields must produce blank output.
    _replace_labeled_block(document, r"^Class and category of danger", r"^2\.2\s+Label elements", fields.get("classification", ""))
    _replace_labeled_block(document, r"^Hazard statements", r"^Supplemental Information", fields.get("hazard_statements", ""))
    _replace_labeled_block(document, r"^Supplemental Information", r"^Precautionary statements", fields.get("supplemental_information", ""))
    _replace_labeled_block(document, r"^Precautionary statements", r"^Pictograms", fields.get("precautionary_statements", ""))
    _replace_labeled_block(document, r"^Pictograms", r"^Other hazards", fields.get("pictograms", ""), remove_drawings=True)
    _replace_labeled_block(document, r"^Other hazards", r"^Section 3", fields.get("other_hazards", ""))
    _replace_section_body(document, r"^4\.2\s+Most important symptoms", r"^4\.3\s+", fields.get("hazard_statements", ""))
    _replace_section_body(document, r"^7\.2\s+Conditions for safe storage", r"^7\.3\s+", fields.get("storage_condition", ""))
    _replace_section_body(document, r"^12\.1\s+Toxicity", r"^12\.2\s+")
    _replace_section_body(document, r"^14\.5\s+Environmental hazards", r"^14\.6\s+")
    _replace_section_body(document, r"^14\.6\s+Special precautions", r"^14\.7\s+")
    _replace_section_body(document, r"^14\.7\s+Maritime transport", r"^Section 15")
    _replace_section_body(document, r"^15\.1\s+Safety, health and environmental", r"^15\.2\s+")
    _replace_section_body(document, r"^15\.2\s+Chemical Safety Assessment", r"^Section 16")
    if not clean_issue_value(fields.get("hazard_statements")):
        _replace_section_body(document, r"^11\.1\s+Information on hazard classes", r"^Information about hazardous ingredients")
    for heading, following in (
        (r"^12\.2\s+Persistence and degradability", r"^12\.3\s+"),
        (r"^12\.3\s+Bio accumulative potential", r"^12\.4\s+"),
        (r"^12\.4\s+Mobility in soil", r"^12\.5\s+"),
        (r"^12\.5\s+Results of PBT", r"^12\.6\s+"),
        (r"^12\.6\s+Endocrine disrupting properties", r"^12\.7\s+"),
        (r"^12\.7\s+Other adverse effects", r"^Section 13"),
        (r"^Key to revisions", r"^Key to abbreviations"),
    ):
        _replace_section_body(document, heading, following)

    composition_tables = [
        table for table in document.tables
        if table.rows and "cas" in normalise(" ".join(c.text for c in table.rows[0].cells))
        and "%" in " ".join(c.text for c in table.rows[0].cells)
        and len(table.rows[0].cells) == 6
    ]
    if not composition_tables:
        _remove_output_placeholders(document)
        return
    composition = composition_tables[0]
    continuation = composition_tables[1] if len(composition_tables) > 1 else None
    primary_template = deepcopy(composition.rows[1]._tr if len(composition.rows) > 1 else composition.rows[0]._tr)
    continuation_template = deepcopy(continuation.rows[1]._tr if continuation is not None and len(continuation.rows) > 1 else primary_template)
    # The supplied Aromazen master is laid out for nine rows on page 2 and a
    # continuation table on page 3. Preserve that pagination instead of
    # pushing the manually positioned headers and footers onto extra pages.
    _populate_composition_table(composition, primary_template, ingredients[:9])
    if continuation is not None:
        _populate_composition_table(continuation, continuation_template, ingredients[9:])

    # Product-specific toxicology: replace all sample-product rows. A single
    # reviewed free-text value is kept in the first result column.
    toxicity = next((table for table in document.tables if table.rows and "ld50 ate oral" in normalise(" ".join(c.text for c in table.rows[0].cells))), None)
    if toxicity is not None:
        template_row = deepcopy(toxicity.rows[1]._tr if len(toxicity.rows) > 1 else toxicity.rows[0]._tr)
        while len(toxicity.rows) > 1:
            toxicity._tbl.remove(toxicity.rows[-1]._tr)
        for item in ingredients:
            value = clean_issue_value(item.get("toxicology"))
            if not value:
                continue
            toxicity._tbl.append(deepcopy(template_row)); cells = toxicity.rows[-1].cells
            values = [item.get("name"), item.get("cas"), item.get("ec"), value, "", "", ""]
            for index, cell_value in enumerate(values):
                if index < len(cells):
                    cells[index].text = clean_issue_value(cell_value)

    # Transport classification is a mixture-level decision. Never retain the
    # example product's UN number, proper shipping name, class or packing group.
    transport = next((table for table in document.tables if table.rows and "un proper shipping name" in normalise(" ".join(c.text for c in table.rows[0].cells))), None)
    if transport is not None:
        for row in transport.rows[1:]:
            for cell in row.cells[1:]:
                cell.text = ""

    # Clear any sample row in the generic ingredient/value table.
    generic = next((table for table in document.tables if table.rows and normalise(" ".join(c.text for c in table.rows[0].cells)) == "ingredient cas ec description value"), None)
    if generic is not None:
        for row in generic.rows[1:]:
            for cell in row.cells:
                cell.text = ""

    _remove_output_placeholders(document)
    _remove_duplicate_particle_characteristics(document)


def generate_regulatory_docx(template: Path, output: Path, document_type: str, product: str, code: str, fields: dict[str, str], ingredients: list[dict]) -> None:
    if document_type not in DOCUMENT_TYPES:
        raise ValueError("Unsupported regulatory document type.")
    document = Document(template)
    _replace_product(document, clean_issue_value(product), clean_issue_value(code))
    if document_type == "sds":
        _fill_sds(document, fields, ingredients)
    elif document_type == "ifra_certificate":
        _tighten_ifra_certificate(document)
    elif document_type == "ifra_amendment":
        _fill_catalog_tables(document, ingredients, "NIL")
    elif document_type == "allergen_report":
        _fill_catalog_tables(document, ingredients, "-")
    elif document_type == "reach_declaration":
        _fill_catalog_tables(document, ingredients, "NIL")
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
