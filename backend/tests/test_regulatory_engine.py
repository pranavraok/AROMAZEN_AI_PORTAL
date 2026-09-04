from io import BytesIO
from pathlib import Path

from docx import Document
from openpyxl import Workbook

from app.modules.regulatory.engine import (
    clean_issue_value,
    extract_coa_identity,
    extract_coa_properties,
    generate_regulatory_docx,
    parse_regulatory_excel,
)


TEMPLATES = Path(__file__).resolve().parents[1] / "app" / "templates" / "regulatory"


def _workbook() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["Name", "CEDAR AND SAGE"])
    sheet.append(["Code", "FS 12388"])
    sheet.append(["SL NO", "RAWMATERIALS", "%"])
    sheet.append([1, "ISO E SUPER", 10])
    sheet.append([2, "LINALOOL", 4.5])
    stream = BytesIO()
    workbook.save(stream)
    return stream.getvalue()


def _text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    parts.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    for section in document.sections:
        parts.extend(paragraph.text for paragraph in section.footer.paragraphs)
        parts.extend(
            cell.text
            for table in section.footer.tables
            for row in table.rows
            for cell in row.cells
        )
    return "\n".join(parts)


def test_formula_and_coa_extraction() -> None:
    product, code, ingredients = parse_regulatory_excel(_workbook())
    assert (product, code) == ("CEDAR AND SAGE", "FS 12388")
    assert ingredients[0]["name"] == "ISO E SUPER"
    assert ingredients[1]["concentration"] == "4.5"
    assert extract_coa_properties("Appearance: Clear liquid\nOdour: Woody\nFlash Point: 82 °C") == {
        "appearance": "Clear liquid", "odour": "Woody", "flash_point": "82 °C"
    }
    assert extract_coa_properties("Flash Point | For record | 990C")["flash_point"] == "99°C"
    assert clean_issue_value("AI suggested - verify") == ""
    assert clean_issue_value("N/A") == ""
    assert clean_issue_value("Leave blank if unavailable") == ""


def test_coa_property_extraction_stops_before_approval_fields() -> None:
    text = "Storage Condition: Store in a cool place\nTested By: Sneha\nChecked By: Rakshanda"
    assert extract_coa_properties(text)["storage_condition"] == "Store in a cool place"


def test_coa_identity_extraction() -> None:
    text = "Date: 27-08-2026\nName of the Product : PEARL\nProduct Code : FPM 10691\nBatch Number: X1"
    assert extract_coa_identity(text) == {"product_name": "PEARL", "product_code": "FPM 10691"}


def test_all_regulatory_documents_are_generated_without_internal_labels(tmp_path: Path) -> None:
    ingredients = [
        {"name": "ISO E SUPER", "cas": "54464-57-2", "ec": "259-174-3", "concentration": "10", "classification": "Skin Irrit. 2: H315", "toxicology": "LD50 oral 5000 mg/kg"},
        {"name": "LINALOOL", "cas": "78-70-6", "ec": "201-134-4", "concentration": "4.5", "classification": "Skin Sens. 1B: H317", "allergen_identity": "Linalool"},
    ]
    fields = {
        "appearance": "Clear liquid", "odour": "Woody", "flash_point": "82 °C",
        "classification": "Skin Irrit. 2: H315", "supplemental_information": "EUH208",
        "other_hazards": "", "version": "0.0", "revision_date": "28-08-2026",
    }
    files = {
        "sds": "sds.docx", "ifra_certificate": "ifra-certificate.docx",
        "ifra_amendment": "ifra-amendment.docx", "allergen_report": "allergen-report.docx",
        "reach_declaration": "reach-declaration.docx",
    }
    for document_type, filename in files.items():
        output = tmp_path / filename
        generate_regulatory_docx(TEMPLATES / filename, output, document_type, "CEDAR AND SAGE", "FS 12388", fields, ingredients)
        text = _text(output).lower()
        assert "ai generated" not in text
        assert "ai suggested" not in text
        assert "review required" not in text
        assert "source url" not in text
        assert "chandan" not in text
        assert "13-07-2026" not in text

    assert "PRODUCT NAME: CEDAR AND SAGE FS 12388" in _text(tmp_path / "ifra-certificate.docx")
    assert "PRODUCTNAME:CEDARANDSAGEFS12388" in _text(tmp_path / "allergen-report.docx").replace(" ", "")
    assert "28-08-2026" in _text(tmp_path / "ifra-certificate.docx")
    assert "28-08-2026" in _text(tmp_path / "ifra-amendment.docx")

    sds_text = _text(tmp_path / "sds.docx")
    assert "PTBCHA" not in sds_text
    assert "GALAXOLIDE" not in sds_text
    assert "ISO E SUPER" in sds_text
    assert "CHANDAN" not in sds_text
    assert "H412, Harmful to aquatic life with long lasting effects" not in sds_text
    assert "Product identifier: CEDAR AND SAGE FS 12388" in sds_text
    assert "1.1 Product Identifier CEDAR" not in sds_text
    assert "Skin Irrit. 2: H315" in sds_text
    assert "EUH208" in sds_text
    lowered_sds = sds_text.lower()
    assert "not determined" not in lowered_sds
    assert "not available" not in lowered_sds
    assert "no available" not in lowered_sds
    assert "none available" not in lowered_sds
    assert "not applicable" not in lowered_sds
    assert "13-07-2026" not in sds_text
    assert lowered_sds.count("particle characteristics:") == 1

    sds = Document(tmp_path / "sds.docx")
    composition_tables = [
        table for table in sds.tables
        if table.rows and "cas" in " ".join(cell.text.lower() for cell in table.rows[0].cells)
        and "%" in " ".join(cell.text for cell in table.rows[0].cells)
        and len(table.rows[0].cells) == 6
    ]
    assert len(composition_tables[0].rows) == 3
    assert len(composition_tables[1].rows) == 1
    first_value_run = composition_tables[0].rows[1].cells[0].paragraphs[0].runs[0]
    assert first_value_run.font.size.pt == 8.5
