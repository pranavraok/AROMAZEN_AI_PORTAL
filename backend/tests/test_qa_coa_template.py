from pathlib import Path

from docx import Document

from app.modules.document_generator.engine import coa_parameter_rows, field_schema, generate_docx


ASSET = Path(__file__).resolve().parents[1] / "app" / "assets" / "qa" / "coa-master.docx"


def test_seeded_qa_coa_master_maps_fields_and_editable_rows(tmp_path: Path) -> None:
    schema = field_schema("coa", ASSET)
    assert {field["key"] for field in schema} >= {
        "date", "product_name", "product_code", "batch_number",
        "manufacturing_date", "expiry_date", "tested_by", "checked_by",
    }
    rows = coa_parameter_rows(ASSET)
    assert [row["parameter"] for row in rows] == [
        "Appearance", "Odour", "Specific Gravity", "Flash Point", "Fire Point", "Refractive Index",
    ]

    output = tmp_path / "generated-coa.docx"
    warnings = generate_docx(
        ASSET,
        output,
        "coa",
        {
            "date": "3 September 2026",
            "product_name": "Rose Absolute",
            "product_code": "RA 101",
            "batch_number": "B25",
            "manufacturing_date": "3 September 2026",
            "expiry_date": "3 September 2028",
        },
        [
            {"parameter": "Acid Value", "specification": "Maximum 2.0", "result": "1.4"},
            {"parameter": "Custom Appearance", "specification": "Clear liquid", "result": "Passes"},
        ],
    )

    assert warnings == []
    generated = Document(output)
    text = "\n".join(paragraph.text for paragraph in generated.paragraphs)
    assert "Date: 03-09-2026" in text
    assert "Name of the Product\t: Rose Absolute" in text
    assert "Product Code\t: RA 101" in text
    assert len(generated.tables[0].rows) == 3
    assert [cell.text for cell in generated.tables[0].rows[1].cells] == ["Acid Value", "Maximum 2.0", "1.4"]
    assert [cell.text for cell in generated.tables[0].rows[2].cells] == ["Custom Appearance", "Clear liquid", "Passes"]
